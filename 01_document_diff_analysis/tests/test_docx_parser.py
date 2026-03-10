"""Tests for DOCX parser behavior."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from docdiff.model import HeadingBlock, ListBlock, TableBlock
from docdiff.parsers import parse_docx_file


def _build_docx_fixture(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Release Notes", level=1)
    doc.add_paragraph("This document captures parser and diff milestones.")

    item_one = doc.add_paragraph("Normalize all textual input")
    item_one.style = "List Bullet"
    item_two = doc.add_paragraph("Preserve deterministic block identifiers")
    item_two.style = "List Bullet"

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Module"
    table.rows[0].cells[1].text = "Status"
    table.rows[1].cells[0].text = "docx"
    table.rows[1].cells[1].text = "ready"

    doc.add_paragraph("Closing paragraph for additional detail.")
    doc.save(str(path))


def test_docx_parser_handles_mixed_content(tmp_path: Path) -> None:
    """DOCX parser should preserve heading/paragraph/list/table order."""
    fixture = tmp_path / "sample.docx"
    _build_docx_fixture(fixture)

    parsed = parse_docx_file(fixture)

    assert [block.block_type for block in parsed.blocks] == [
        "heading",
        "paragraph",
        "list",
        "table",
        "paragraph",
    ]
    heading_block = parsed.blocks[0]
    assert isinstance(heading_block, HeadingBlock)
    assert heading_block.text == "Release Notes"
    list_block = parsed.blocks[2]
    assert isinstance(list_block, ListBlock)
    assert list_block.items[0] == "Normalize all textual input"
    table_block = parsed.blocks[3]
    assert isinstance(table_block, TableBlock)
    assert table_block.header == ["Module", "Status"]
    assert table_block.rows == [["docx", "ready"]]
